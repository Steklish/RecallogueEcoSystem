# app/ingest.py
from __future__ import annotations
import os, re
from typing import List
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
import chardet

# --------- простая, без NLTK, токенизация/чанкинг ---------

# Разделяем по . ! ? … или ... , с возможными завершающими кавычками/скобками
# и допускаем кавычки/скобки/пробелы перед началом следующего предложения.
_SENT_SPLIT = re.compile(
    r"""                # конец предложения:
        (?:
            (?<!\b[А-ЯA-Z]\.)   # простая защита от инициалов "И." (эвристика)
            (?<!\b[А-ЯA-Z]\.[А-ЯA-Z]\.)  # "И.О."
            (?<!\bт\.д) (?<!\bт\.п) (?<!\bи\.т\.д) (?<!\bи\.т\.п)  # популярные сокращения
            (?<=[\.\!\?])       # обычные . ! ?
            | (?<=\u2026)       # …
            | (?<=\.\.\.)       # ...
        )
        [\"»”)\]]*              # завершающие кавычки/скобки
        \s+                     # пробелы/перевод строки
        (?=[^\s])               # затем что-то «начинается»
    """,
    flags=re.U | re.X
)

_WS = re.compile(r"[ \t\u00A0\u2000-\u200B\u202F\u205F\u3000]+")  # добавлены тонкие/узкие/идеографические пробелы

_LOW = r"[a-zа-яё]"

_HYPHENS = r"[\-\u00AD\u2010\u2011]"

def _split_sentences(text: str) -> List[str]:
    # если предложений мало — режем по переносам
    if "\n" in text and len(text) < 2000:
        parts = [p.strip() for p in text.splitlines() if p.strip()]
        return parts if len(parts) > 1 else [text.strip()]
    # обычный режим
    parts = _SENT_SPLIT.split(text)
    parts = [p.strip() for p in parts if p and p.strip()]
    return parts if parts else [text.strip()]

def extract_text_from_file(path: str, mime: str | None = None) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx" or (mime and "wordprocessingml" in mime):
        return extract_docx(path)
    if ext in [".html", ".htm"]:
        return extract_html(path)
    # всё остальное — как текст
    return _read_text_best_effort(path)


def _dehyphenate_lines(text: str) -> str:
    """
    Сшивает слова, разорванные переносом строки (только если по обе стороны строчные буквы),
    и удаляет мягкие дефисы. Пример: "проек-\nтирование" -> "проектирование".
    """
    # Мягкий дефис (discretionary hyphen) удаляем везде — безопасно
    text = text.replace("\u00AD", "")
    # Удаляем дефис на конце строки, если вокруг строчные буквы (латиница/кириллица)
    text = re.sub(fr"({_LOW}){_HYPHENS}\s*\n({_LOW})", r"\1\2", text, flags=re.U)
    return text


def extract_pdf(path: str) -> str:
    """
    Извлекает текст из PDF постранично, затем убирает переносы слов и нормализует пробелы/пустые строки.
    Ожидает, что в модуле определена функция normalize_text(text: str) -> str.
    """
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                parts.append(text)

    text = "\n".join(parts)
    text = _dehyphenate_lines(text)
    return normalize_text(text)


def _read_text_best_effort(path: str) -> str:
    with open(path, "rb") as f:
        raw_data = f.read()

    # 1. Try UTF-8 first, as it's most common
    try:
        return raw_data.decode("utf-8")
    except UnicodeDecodeError:
        pass

    # 2. If UTF-8 fails, use chardet
    detection = chardet.detect(raw_data)
    encoding = detection.get("encoding")
    
    if encoding:
        try:
            return raw_data.decode(encoding)
        except (UnicodeDecodeError, TypeError):
            pass # It failed, so we'll try other encodings

    # 3. Try common fallbacks
    for enc in ("cp1251", "latin-1"):
        try:
            return raw_data.decode(enc)
        except (UnicodeDecodeError, TypeError):
            continue
    
    # 4. Last resort
    return raw_data.decode("utf-8", errors="ignore")



def extract_docx(path: str) -> str:
    """
    Полный проход: абзацы + таблицы (в т.ч. объединённые ячейки).
    Таблицы представляем строками формата: "col1 | col2 | col3".
    """
    doc = DocxDocument(path)
    out: List[str] = []

    # 1) абзацы
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)

    # 2) таблицы
    for table in doc.tables:
        for row in table.rows:
            cells = []
            # row.cells уже учитывает merges; собираем все параграфы ячейки
            for cell in row.cells:
                ct = "\n".join((para.text or "").strip() for para in cell.paragraphs if (para.text or "").strip())
                ct = ct.strip()
                cells.append(ct)
            # уберём пустые по краям, но сохраним структуру
            # (если вся строка пустая — пропустим)
            if any(cells):
                line = " | ".join(cells)
                out.append(line)

    # 3) сжатие пустых строк/пробелов
    text = "\n".join(out)
    return normalize_text(text)

def extract_html(path: str) -> str:
    html = _read_text_best_effort(path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    text = soup.get_text("\n")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)

def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS.sub(" ", text)
    lines = [l.strip() for l in text.split("\n")]
    out, blank = [], False
    for l in lines:
        if not l:
            if not blank:
                out.append(""); blank = True
        else:
            out.append(l); blank = False
    return "\n".join(out).strip()

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> List[str]:
    # chunk_size и overlap считаем в словах
    sents = _split_sentences(text)
    chunks, cur, cur_len = [], [], 0
    for s in sents:
        slen = len(s.split())
        if cur and cur_len + slen > chunk_size:
            joined = " ".join(cur).strip()
            if joined:
                chunks.append(joined)
            if overlap > 0:
                # возьмём хвост из последнего предложения (или двух), а не по словам
                tail_sents = _split_sentences(joined)
                tail = " ".join(tail_sents[-2:]) if len(tail_sents) >= 2 else (tail_sents[-1] if tail_sents else "")
                cur, cur_len = ([tail] if tail else []), len(tail.split()) if tail else 0
            else:
                cur, cur_len = [], 0
        cur.append(s); cur_len += slen
    if cur:
        joined = " ".join(cur).strip()
        if joined:
            chunks.append(joined)
    # фильтр совсем коротких
    return [c for c in chunks if len(c.split()) >= 5]